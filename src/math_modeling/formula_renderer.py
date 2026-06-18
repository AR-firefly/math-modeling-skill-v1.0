"""
公式渲染器：LaTeX → PNG → 嵌入文档
"""
import os
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt


def render_latex(latex_str, width_inches=4.5, dpi=150, fontsize=12,
                 save_path=None):
    """
    将 LaTeX 公式渲染为 PNG 图片。

    Args:
        latex_str: LaTeX 公式字符串（不含 $$ 分隔符）
        width_inches: 图片宽度（英寸）
        dpi: 输出分辨率
        fontsize: 字体大小
        save_path: 保存路径，None 则返回 matplotlib figure

    Returns:
        如果 save_path 指定则返回路径，否则返回 (fig, ax)
    """
    fig, ax = plt.subplots(figsize=(width_inches, 0.5))
    ax.axis('off')
    try:
        ax.text(0.5, 0.5, f'${latex_str}$', transform=ax.transAxes,
                fontsize=fontsize, ha='center', va='center')
    except Exception:
        ax.text(0.5, 0.5, latex_str, transform=ax.transAxes,
                fontsize=fontsize - 1, ha='center', va='center',
                family='monospace')

    if save_path:
        fig.savefig(save_path, dpi=dpi, bbox_inches='tight', pad_inches=0.1)
        plt.close(fig)
        return save_path
    return fig, ax


def render_and_embed(doc, latex_str, width_inches=4.5, temp_dir=None):
    """
    渲染 LaTeX 公式为图片并嵌入 python-docx 段落。

    Args:
        doc: python-docx Document 对象
        latex_str: LaTeX 公式
        width_inches: 图片宽度
        temp_dir: 临时目录，None 则使用当前目录
    """
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    temp_dir = temp_dir or os.getcwd()
    path = os.path.join(temp_dir, '_formula_temp.png')
    render_latex(latex_str, width_inches=width_inches, save_path=path)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Inches(0)
    run = p.add_run()
    run.add_picture(path, width=Inches(width_inches))
    return p
