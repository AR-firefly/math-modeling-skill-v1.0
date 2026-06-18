"""
三线表生成器：专为数学建模论文设计的表格构建工具
"""
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


class TableBuilder:
    """三线表构建器"""

    def __init__(self, doc, headers, col_widths=None,
                 font_name='宋体', font_size=10, header_bg='D9D9D9'):
        """
        Args:
            doc: python-docx Document 对象
            headers: 表头列表
            col_widths: 列宽列表（可选）
            font_name: 字体
            font_size: 字号
            header_bg: 表头背景色（RGB hex）
        """
        self.doc = doc
        self.headers = headers
        self.col_widths = col_widths
        self.font_name = font_name
        self.font_size = font_size
        self.header_bg = header_bg
        self.rows = []

    def add_row(self, row_data):
        """添加数据行"""
        self.rows.append(row_data)
        return self

    def add_rows(self, rows_data):
        """批量添加数据行"""
        self.rows.extend(rows_data)
        return self

    def build(self, caption=None, caption_font_size=10):
        """
        构建表格并插入文档。

        Args:
            caption: 表题文字
            caption_font_size: 表题字号
        """
        table = self.doc.add_table(rows=1 + len(self.rows),
                                   cols=len(self.headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'

        # 表头
        for j, h in enumerate(self.headers):
            cell = table.rows[0].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(h)
            r.bold = True
            r.font.size = Pt(self.font_size)
            r.font.name = self.font_name
            r._element.rPr.rFonts.set(qn('w:eastAsia'), self.font_name)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            shading = parse_xml(
                f'<w:shd {nsdecls("w")} w:fill="{self.header_bg}"/>')
            cell._element.get_or_add_tcPr().append(shading)

        # 数据行
        for i, row in enumerate(self.rows):
            for j, val in enumerate(row):
                cell = table.rows[i + 1].cells[j]
                cell.text = ''
                p = cell.paragraphs[0]
                r = p.add_run(str(val))
                r.font.size = Pt(self.font_size)
                r.font.name = self.font_name
                r._element.rPr.rFonts.set(qn('w:eastAsia'), self.font_name)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if caption:
            cap = self.doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.first_line_indent = Cm(0)
            r = cap.add_run(caption)
            r.font.size = Pt(caption_font_size)
            r.font.name = self.font_name
            r._element.rPr.rFonts.set(qn('w:eastAsia'), self.font_name)

        return table
