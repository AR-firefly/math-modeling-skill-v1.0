"""
论文生成器：参数化的数学建模竞赛论文生成框架

使用方式：
    from math_modeling import PaperConfig, PaperGenerator

    config = PaperConfig(
        school_name="示例大学",
        title="AGV协同调度与能耗优化问题研究",
    )
    gen = PaperGenerator(config)
    gen.set_results(results_dict)    # 传入求解结果（可选）
    gen.build()                      # 构建论文
    gen.save("output.docx")          # 保存
"""
import os
import json
import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

try:
    from .table_builder import TableBuilder
    from .formula_renderer import render_latex, render_and_embed
except ImportError:
    from table_builder import TableBuilder
    from formula_renderer import render_latex, render_and_embed


class PaperConfig:
    """论文配置 —— 所有学校/竞赛特定信息参数化"""

    def __init__(self,
                 school_name="示例大学",
                 title="{{论文标题}}",
                 year="2026",
                 team_number="________________（请填写）",
                 members=None,
                 advisor="________________（请填写）",
                 body_font="宋体",
                 heading_font="黑体",
                 body_font_size=12,
                 h1_size=16,
                 h2_size=14,
                 h3_size=13,
                 top_margin_cm=2.54,
                 bottom_margin_cm=2.54,
                 left_margin_cm=3.18,
                 right_margin_cm=3.18,
                 line_spacing=1.5,
                 title_font_size=18):
        self.school_name = school_name
        self.title = title
        self.year = year
        self.team_number = team_number
        self.members = members or ["________________"] * 3
        self.advisor = advisor
        self.body_font = body_font
        self.heading_font = heading_font
        self.body_font_size = body_font_size
        self.h1_size = h1_size
        self.h2_size = h2_size
        self.h3_size = h3_size
        self.top_margin_cm = top_margin_cm
        self.bottom_margin_cm = bottom_margin_cm
        self.left_margin_cm = left_margin_cm
        self.right_margin_cm = right_margin_cm
        self.line_spacing = line_spacing
        self.title_font_size = title_font_size


class PaperGenerator:
    """论文生成器"""

    def __init__(self, config: PaperConfig, result_dir=None):
        self.config = config
        self.result_dir = result_dir or os.getcwd()
        self.results = None  # set_results() 注入求解结果
        self.doc = Document()
        self._setup_page()
        self._setup_styles()

    # ── 结果注入 ──
    def set_results(self, results_dict):
        """
        注入求解结果，论文生成时自动填入真实数值。

        results_dict 结构：

        results = {
            'config': {'N_STATIONS': 6, 'N_AGVS': 2, 'N_TASKS': 10, ...},
            'phase1': {'dp_distance': 241.51, 'dp_route': [0,4,6,2,3,5,1,0]},
            'phase2': {'makespan': 267.1, 'total_trips': 6, 'penalty': 0.0,
                       'agv_trips': {0: [[trip], [trip]], 1: [[trip]]}},
            'phase3': {'knee_time': 271.0, 'knee_energy': 277.6,
                       'pareto_points': 100, 'min_time': 267.1, 'min_energy': 277.6},
            'phase4': {'best_station': 1, 'best_x': 35, 'best_y': 60,
                       'mksp_improve': 9.3, 'energy_improve': 17.6},
        }
        """
        self.results = results_dict
        return self

    def from_results_json(self, json_path):
        """从 results.json 文件读取结果"""
        with open(json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        # 尝试读取同目录的 config info
        config_path = os.path.join(os.path.dirname(json_path), 'config_summary.json')
        if os.path.exists(config_path):
            with open(config_path, 'r') as cf:
                data['config'] = json.load(cf)
        self.set_results(data)
        return self

    def _r(self, section, key, fmt=None):
        """从 results 中安全取值"""
        if self.results is None:
            return None
        val = self.results.get(section, {}).get(key)
        if val is None:
            return None
        if fmt == '.1f':
            return f'{val:.1f}'
        if fmt == '.2f':
            return f'{val:.2f}'
        if fmt == '.0f':
            return f'{val:.0f}'
        return val

    def _fmt_route(self, route):
        return ' → '.join(map(str, route))

    # ── 页面设置 ──
    def _setup_page(self):
        section = self.doc.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(self.config.top_margin_cm)
        section.bottom_margin = Cm(self.config.bottom_margin_cm)
        section.left_margin = Cm(self.config.left_margin_cm)
        section.right_margin = Cm(self.config.right_margin_cm)

    # ── 样式 ──
    def _setup_styles(self):
        style = self.doc.styles['Normal']
        style.font.name = self.config.body_font
        style.font.size = Pt(self.config.body_font_size)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), self.config.body_font)
        style.paragraph_format.line_spacing = self.config.line_spacing
        style.paragraph_format.first_line_indent = Cm(0.74)

        for i, size in [(1, self.config.h1_size),
                        (2, self.config.h2_size),
                        (3, self.config.h3_size)]:
            h = self.doc.styles[f'Heading {i}']
            h.font.name = self.config.heading_font
            h.element.rPr.rFonts.set(qn('w:eastAsia'), self.config.heading_font)
            h.font.color.rgb = RGBColor(0, 0, 0)
            h.font.size = Pt(size)
            h.paragraph_format.first_line_indent = Cm(0)

    # ── 段落辅助 ──
    def add_para(self, text, bold=False, font_size=None, alignment=None,
                 indent=True, font_name=None):
        p = self.doc.add_paragraph()
        if not indent:
            p.paragraph_format.first_line_indent = Cm(0)
        if alignment is not None:
            p.alignment = alignment
        run = p.add_run(text)
        run.bold = bold
        fs = font_size or self.config.body_font_size
        run.font.size = Pt(fs)
        fn = font_name or self.config.body_font
        run.font.name = fn
        run._element.rPr.rFonts.set(qn('w:eastAsia'), fn)
        return p

    def add_bold_normal(self, bold_text, normal_text, font_size=None):
        fs = font_size or self.config.body_font_size
        p = self.doc.add_paragraph()
        for text, bold in [(bold_text, True), (normal_text, False)]:
            r = p.add_run(text)
            r.bold = bold
            r.font.size = Pt(fs)
            r.font.name = self.config.body_font
            r._element.rPr.rFonts.set(qn('w:eastAsia'), self.config.body_font)
        return p

    def add_heading(self, text, level=1):
        h = self.doc.add_heading(text, level=level)
        h.paragraph_format.first_line_indent = Cm(0)
        return h

    def add_image(self, path, caption, width_inches=5.5):
        if not os.path.exists(path):
            self.add_para(f'[图片缺失: {path}]', font_size=10)
            return
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run()
        run.add_picture(path, width=Inches(width_inches))
        cap = self.doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.first_line_indent = Cm(0)
        r = cap.add_run(caption)
        r.font.size = Pt(10)
        r.font.name = self.config.body_font
        r._element.rPr.rFonts.set(qn('w:eastAsia'), self.config.body_font)

    def add_table(self, headers, rows, caption=None):
        tb = TableBuilder(self.doc, headers,
                          font_name=self.config.body_font)
        tb.add_rows(rows)
        tb.build(caption=caption)

    def add_formula(self, latex_str, width_inches=4.5):
        render_and_embed(self.doc, latex_str, width_inches=width_inches)

    def page_break(self):
        self.doc.add_page_break()

    # ═══════════════════════════════════════
    # 内容模块
    # ═══════════════════════════════════════

    def write_cover(self):
        """封面（参数化，无学校特定格式）"""
        c = self.config
        for _ in range(4):
            self.doc.add_paragraph()

        self.add_para(c.school_name, font_size=22,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER,
                      indent=False, font_name=c.heading_font)
        self.add_para(f'{c.year}年参赛论文', font_size=18,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER,
                      indent=False, font_name=c.heading_font)

        for _ in range(3):
            self.doc.add_paragraph()

        info = [
            ('参赛题号：', '______'),
            ('参赛队号：', c.team_number),
            ('队员1（签名）：', c.members[0]),
            ('队员2（签名）：', c.members[1] if len(c.members) > 1 else '________________'),
            ('队员3（签名）：', c.members[2] if len(c.members) > 2 else '________________'),
            ('指导教师（签名）：', c.advisor),
            ('参赛日期：', f'{c.year}年____月____日'),
        ]
        for label, value in info:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            r1 = p.add_run(label)
            r1.font.size = Pt(14)
            r1.font.name = c.body_font
            r1._element.rPr.rFonts.set(qn('w:eastAsia'), c.body_font)
            r2 = p.add_run(value)
            r2.font.size = Pt(14)
            r2.font.name = c.body_font
            r2._element.rPr.rFonts.set(qn('w:eastAsia'), c.body_font)
            r2.underline = True

        for _ in range(4):
            self.doc.add_paragraph()
        self.add_para('我们承诺：严格遵守竞赛规则，独立完成论文，尊重知识产权。',
                      alignment=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
        self.add_para('（此页为封面，请打印后签字装订）', font_size=10,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
        self.page_break()

    def write_abstract(self):
        """摘要（总分总结构，内容为示例）"""
        c = self.config
        self.add_para(c.title, font_size=c.title_font_size,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER,
                      indent=False, bold=True, font_name=c.heading_font)
        self.doc.add_paragraph()
        self.add_heading('摘要', level=1)

        # 总开头
        self.add_para(
            '本文针对现代制造车间中的AGV调度与能耗优化问题，'
            '建立从单AGV路径规划到多目标协同优化的递进式模型体系。'
            '模型综合考虑了行驶距离、时间窗约束、载重约束和能耗因素，'
            '为车间AGV管理提供完整的决策支持框架。'
            '（注：本文内容为示例，实际使用时请替换为具体问题描述和结果。）')

        # ── 分：各问一段（有结果则填入，无结果则通用描述） ──
        p1_dist = self._r('phase1', 'dp_distance', '.2f')
        p1_route = self._r('phase1', 'dp_route')
        p2_makespan = self._r('phase2', 'makespan', '.1f')
        p2_trips = self._r('phase2', 'total_trips', '.0f')
        p2_penalty = self._r('phase2', 'penalty', '.1f')
        p3_knee_t = self._r('phase3', 'knee_time', '.1f')
        p3_knee_e = self._r('phase3', 'knee_energy', '.1f')
        p3_count = self._r('phase3', 'pareto_points', '.0f')
        p4_station = self._r('phase4', 'best_station')
        p4_mksp = self._r('phase4', 'mksp_improve', '.1f')
        p4_energy = self._r('phase4', 'energy_improve', '.1f')
        p4_x = self._r('phase4', 'best_x', '.0f')
        p4_y = self._r('phase4', 'best_y', '.0f')

        if p1_dist:
            route_str = self._fmt_route(p1_route) if p1_route else ''
            self.add_para(
                f'针对Phase1（单AGV路径规划），建立TSP模型，采用Held-Karp动态规划精确求解。'
                f'得到最优路径长度为{p1_dist}m，路线为{route_str}。'
                f'通过最近邻贪心基线对比验证了结果正确性。')
        else:
            self.add_para(
                '针对Phase1（单AGV路径规划），建立TSP模型，采用Held-Karp动态规划精确求解。'
                '算法在指数级搜索空间中高效找到全局最优路径。')

        if p2_makespan:
            self.add_para(
                f'针对Phase2（多AGV协同调度），建立VRPTW扩展模型。'
                f'设计增强遗传算法，融合FFD装箱、LPT负载均衡和2-opt局部搜索，'
                f'得到最优调度方案：makespan={p2_makespan}s，共{p2_trips}趟配送，'
                f'超时惩罚为{p2_penalty}。')
        else:
            self.add_para(
                '针对Phase2（多AGV协同调度），建立VRPTW扩展模型。'
                '设计增强遗传算法，融合FFD装箱、LPT负载均衡和2-opt局部搜索，'
                '在满足容量和时间窗约束下最小化最大完成时间。')

        if p3_knee_t:
            self.add_para(
                f'针对Phase3（时间-能耗双目标优化），建立多目标优化模型，'
                f'采用NSGA-II配合VNS局部搜索求解。'
                f'获得{p3_count}个非支配解构成Pareto前沿，'
                f'Knee点方案为：makespan={p3_knee_t}s，总能耗={p3_knee_e}kJ。')
        else:
            self.add_para(
                '针对Phase3（时间-能耗双目标优化），建立多目标优化模型，'
                '采用NSGA-II配合VNS局部搜索求解。'
                '获得分布均匀的Pareto前沿，Knee点提供了最佳折中方案。')

        if p4_station is not None:
            self.add_para(
                f'针对Phase4（布局微调优化），设计两阶段搜索策略。'
                f'结果表明，移动工位{p4_station}至({p4_x},{p4_y})，'
                f'可使makespan改善{p4_mksp}%、总能耗改善{p4_energy}%。')
        else:
            self.add_para(
                '针对Phase4（布局微调优化），设计两阶段搜索策略——'
                'TSP快速筛选缩小候选范围，GA精评确定最优调整方案。')

        # 总结尾
        self.add_para(
            '本文模型考虑了实际生产中的多种约束条件，'
            '经灵敏度分析验证了参数鲁棒性。'
            '整体形成了"路径规划—协同调度—多目标优化—布局改善"的递进式决策框架。')

        self.add_para(
            '【关键词】AGV协同调度；遗传算法；NSGA-II；能耗优化；布局优化',
            indent=False)

        self.page_break()

    def write_restatement(self):
        """问题重述（示例内容，用户替换）"""
        self.add_heading('一、问题重述', level=1)
        self.add_heading('1.1 问题背景', level=2)
        self.add_para(
            '在现代制造车间中，自动导引车（AGV）作为核心物流装备，'
            '承担着物料搬运任务。多AGV的协同调度与能耗优化'
            '直接影响生产效率和运营成本。本示例问题研究'
            '一个典型车间场景中的AGV调度优化问题。')

        self.add_heading('1.2 问题描述', level=2)
        self.add_para(
            '某智能车间包含若干工位和仓储中心，配备多台同质AGV。'
            '需在满足载重和时间约束下完成配送任务。需解决四个递进式的优化问题：')

        phases = [
            'Phase1：单AGV路径规划。不考虑容量和时间窗，求遍历全部工位的最短路径。',
            'Phase2：多AGV协同调度。多台AGV协同完成任务，满足载重和时间窗约束。',
            'Phase3：多目标能耗优化。同时优化时间和能耗，分析权衡关系。',
            'Phase4：车间布局微调。调整工位位置以改善调度效率。',
        ]
        for p in phases:
            self.add_para(p)
        self.add_para(
            '（注：以上为示例问题描述。实际使用时替换为具体竞赛题目内容。）',
            font_size=10)
        self.page_break()

    def write_analysis(self):
        """问题分析"""
        self.add_heading('二、问题分析', level=1)

        self.add_heading('2.1 Phase1分析', level=2)
        self.add_para(
            'Phase1本质为TSP问题。当工位数量较少时，'
            '可采用精确动态规划算法（Held-Karp）在可接受时间内求得全局最优解。'
            '同时以贪心算法作为基线对比。')

        self.add_heading('2.2 Phase2分析', level=2)
        self.add_para(
            'Phase2为VRPTW的多车扩展，属于NP-hard问题。'
            '当任务规模较大时无法精确求解，需采用元启发式方法。'
            'GA具有全局搜索能力强、不依赖梯度信息的特点。')

        self.add_heading('2.3 Phase3分析', level=2)
        self.add_para(
            '时间与能耗存在天然冲突，构成多目标优化问题。'
            '采用Pareto方法可一次运行获得完整权衡信息。'
            'NSGA-II配合局部搜索可在收敛性和分布性间取得平衡。')

        self.add_heading('2.4 Phase4分析', level=2)
        self.add_para(
            '连续空间位置优化搜索空间大，'
            '采用两阶段策略（TSP快速筛选+GA精评）可有效降低计算量。')
        self.page_break()

    def write_assumptions(self):
        """模型假设"""
        self.add_heading('三、模型假设', level=1)
        self.add_para('为简化问题、聚焦核心矛盾，本文做出以下合理假设：')

        assumptions = [
            ('假设1（匀速行驶假设）',
             'AGV保持恒定速度行驶（速度值由配置参数设定），忽略加速减速。'),
            ('假设2（开放平面假设）',
             '车间内任意两点直线可达，距离按欧氏距离计算。'
             '该假设是简化系统的核心假设——若存在障碍物则需引入避障约束。'),
            ('假设3（固定装卸时间假设）',
             '装卸货时间恒定（由配置参数设定），忽略工位忙碌导致的等待。'),
            ('假设4（电量充足假设）',
             'AGV电池足够完成全部任务，中途无需充电。'),
        ]
        for title, content in assumptions:
            self.add_bold_normal(title + '：', content)

        self.add_para(
            '（注：以上假设为常见简化假设，用户应根据具体问题调整。）',
            font_size=10)
        self.page_break()

    def write_symbols(self):
        """符号说明"""
        self.add_heading('四、符号说明', level=1)
        self.add_para('本文使用的主要符号及其说明如下表所示。示例参数值仅供参考。',
                      indent=False)

        headers = ['符号', '说明', '单位']
        rows = [
            ['v', 'AGV行驶速度', 'm/s'],
            ['t₀', '工位装卸货时间', 's'],
            ['Q', '单台AGV最大载重量', 'kg'],
            ['α', 'AGV基础能耗系数', 'kJ/m'],
            ['β', 'AGV载重附加能耗系数', 'kJ/(m·kg)'],
            ['N', '配送任务总数', '—'],
            ['M', 'AGV数量', '—'],
            ['K', '工位数量', '—'],
            ['d_ij', '节点i到j的欧氏距离', 'm'],
            ['x_ij', '0-1决策变量', '—'],
            ['w_k', '第k个任务的货物重量', 'kg'],
            ['T_k', '第k个任务的截止时间', 's'],
            ['δ', '超时宽限期', 's'],
            ['C_max', '最大完成时间（makespan）', 's'],
            ['E_total', '总能耗', 'kJ'],
        ]
        self.add_table(headers, rows, caption='表1  符号说明表（示例参数，用户替换）')
        self.page_break()

    def write_phase1(self):
        """Phase1 模型与求解"""
        self.add_heading('五、Phase1：单AGV路径规划的模型建立与求解', level=1)

        self.add_heading('5.1 模型建立', level=2)
        self.add_para(
            '问题可抽象为旅行商问题（TSP）：从仓库出发，遍历全部工位各一次后返回。')

        self.add_para('（1）决策变量', indent=False)
        self.add_para(
            'x_ij ∈ {0,1}：若AGV从节点i直接行驶到节点j则取1，否则取0。')

        self.add_para('（2）目标函数', indent=False)
        self.add_para('最小化总行驶距离：min Z = Σ_i Σ_j d_ij · x_ij')

        self.add_para('（3）约束条件', indent=False)
        for c in [
            '每个工位恰好访问一次：Σ_i x_ij = 1, ∀j',
            '每个工位恰好离开一次：Σ_j x_ij = 1, ∀i',
            '从仓库出发并返回：Σ_j x_0j = 1, Σ_i x_i0 = 1',
            '子环消除：Σ_{i∈S} Σ_{j∉S} x_ij ≥ 1, ∀非空真子集S',
            '0-1约束：x_ij ∈ {0,1}',
        ]:
            self.add_para(c)

        self.add_heading('5.2 模型求解', level=2)
        self.add_para(
            '采用Held-Karp动态规划精确求解。'
            '状态定义：dp[mask][i]表示从0出发、已访问节点集合为mask、'
            '当前位于节点i的最短距离。转移方程：')
        self.add_formula('dp[mask | (1<<j)][j] = min{ dp[mask][i] + d_ij }')
        self.add_para(
            '最终答案：min_i { dp[(1<<K)-1][i] + d(i,0) }。'
            '时间复杂度O(K²·2^K)。同时以最近邻贪心算法作为基线对比。')

        self.add_heading('5.3 求解结果', level=2)
        d = self._r('phase1', 'dp_distance', '.2f')
        route = self._r('phase1', 'dp_route')
        if d:
            self.add_para(f'Held-Karp动态规划求得最短路径总长度为{d}m。')
            self.add_para(f'最优路线：{self._fmt_route(route)}')
        else:
            self.add_para('（运行求解后填入结果。）')
            self.add_para('最短路径总长度：{{XXX.XX}} m')
            self.add_para('最优路线：{{路线}}')

        img = os.path.join(self.result_dir, 'phase1_path.png')
        self.add_image(img, '图1  单AGV最优路径')

        self.page_break()

    def write_phase2(self):
        """Phase2 模型与求解"""
        self.add_heading('六、Phase2：多AGV协同调度的模型建立与求解', level=1)

        self.add_heading('6.1 模型建立', level=2)
        self.add_para(
            'Phase2需为多台AGV分配配送任务。'
            '每台AGV执行多趟配送，每趟从仓库出发、完成任务后返回。')

        self.add_para('（1）决策变量', indent=False)
        self.add_para('任务排列π、趟次划分T_r、AGV分配a_{rk}。')

        self.add_para('（2）目标函数', indent=False)
        self.add_para('最小化makespan：min C_max = max_k T_k')

        self.add_para('（3）约束条件', indent=False)
        for c in [
            '容量约束：每趟总载重 ≤ Q',
            '任务分配约束：每个任务恰好分配一次',
            '时间窗约束（软约束）：超时产生分段线性惩罚',
        ]:
            self.add_para(c)

        self.add_para('（4）惩罚函数', indent=False)
        self.add_formula('P(t) = 0, t ≤ T + δ')
        self.add_formula('P(t) = c1(t - T - δ), T + δ < t ≤ T + 2δ')
        self.add_formula('P(t) = c1·δ + c2(t - T - 2δ), t > T + 2δ')

        self.add_heading('6.2 模型求解——增强遗传算法', level=2)
        self.add_para(
            '采用GA求解，编码为任务排列。解码过程：'
            'FFD装箱划分趟次，LPT规则分配AGV实现负载均衡。'
            '遗传算子包括PMX交叉、多种变异和2-opt局部搜索。')

        self.add_heading('6.3 求解结果', level=2)
        ms = self._r('phase2', 'makespan', '.1f')
        trips = self._r('phase2', 'total_trips', '.0f')
        pen = self._r('phase2', 'penalty', '.1f')
        agv_trips = self._r('phase2', 'agv_trips') if self.results else None

        if ms:
            self.add_para(
                f'增强遗传算法在有限代数内稳定收敛，得到最优调度方案：')
            self.add_para(f'最大完成时间（makespan）：{ms}s')
            self.add_para(f'总配送趟数：{trips}趟')
            self.add_para(f'总超时惩罚：{pen}（所有任务均在截止时间内完成）')

            if agv_trips:
                self.add_para('各AGV任务分配详见表2。', indent=False)
                hdrs = ['AGV', '趟次', '路线', '载重量(kg)', '耗时(s)']
                rows = []
                for agv_id, trip_list in agv_trips.items():
                    aid = int(agv_id)  # JSON key may be string "0"
                    for j, t in enumerate(trip_list):
                        rows.append([f'AGV{aid+1}', f'T{j+1}',
                                     t.get('route_str', ''), t.get('load', ''),
                                     t.get('time', '')])
                self.add_table(hdrs, rows, caption='表2  最优调度方案')
        else:
            self.add_para('（运行求解后填入结果。）')

        img1 = os.path.join(self.result_dir, 'phase2_gantt.png')
        self.add_image(img1, '图2  多AGV调度甘特图')
        img2 = os.path.join(self.result_dir, 'phase2_paths.png')
        self.add_image(img2, '图3  AGV配送路径图')

        self.page_break()

    def write_phase3(self):
        """Phase3 模型与求解"""
        self.add_heading('七、Phase3：多目标能耗优化的模型建立与求解', level=1)

        self.add_heading('7.1 模型建立', level=2)
        self.add_para(
            '在Phase2基础上新增能耗目标，形成双目标优化问题。')

        self.add_para('（1）目标函数', indent=False)
        self.add_para('目标一：min f₁ = C_max')
        self.add_para('目标二：min f₂ = E_total = Σ (α + β·m_load) × d_seg')
        self.add_para(
            '能耗按逐段累加方式计算——AGV在某路段上的能耗取决于当前载重。'
            '空载返回段m=0，能耗为α×d。')

        self.add_para('（2）Pareto最优', indent=False)
        self.add_para(
            '对于双目标优化，不存在同时使两个目标均最优的单一解。'
            '不被任何解支配的解构成Pareto最优解集。')

        self.add_heading('7.2 模型求解——NSGA-II + VNS', level=2)
        self.add_para(
            '采用NSGA-II算法：非支配排序分层，拥挤度距离保持分布性，'
            '精英保留策略防止最优解丢失。'
            'VNS每5代对前沿中的精英解进行局部强化。')

        self.add_heading('7.3 求解结果', level=2)
        kt = self._r('phase3', 'knee_time', '.1f')
        ke = self._r('phase3', 'knee_energy', '.1f')
        pc = self._r('phase3', 'pareto_points', '.0f')
        mt = self._r('phase3', 'min_time', '.1f')
        me = self._r('phase3', 'min_energy', '.1f')

        if kt:
            self.add_para(
                f'NSGA-II运行后获得包含{pc}个非支配解的Pareto前沿。'
                f'前沿分布表明配送时间与总能耗之间存在明显的权衡关系：')

            self.add_para(
                f'Knee点（最佳折中方案）：makespan={kt}s, 总能耗={ke}kJ。'
                f'该点位于Pareto前沿曲率最大处，是时间与能耗的最佳折中。')

            if mt and me:
                self.add_para(
                    f'极速方案：makespan={mt}s（以更高能耗换取最短时间）。'
                    f'节能方案：总能耗={me}kJ（以稍长时间换取最低能耗）。')
        else:
            self.add_para('（运行求解后填入结果。）')

        img1 = os.path.join(self.result_dir, 'phase3_pareto.png')
        self.add_image(img1, '图4  Pareto前沿')
        img2 = os.path.join(self.result_dir, 'phase3_compare.png')
        self.add_image(img2, '图5  策略对比')

        self.page_break()

    def write_phase4(self):
        """Phase4 模型与求解"""
        self.add_heading('八、Phase4：车间布局微调的模型建立与求解', level=1)

        self.add_heading('8.1 模型建立', level=2)
        self.add_para(
            '允许调整工位位置以改善调度效率。'
            '搜索空间为工位坐标在[-R, R]范围内的连续区域。')

        self.add_heading('8.2 两阶段搜索策略', level=2)
        self.add_para(
            'Phase 1（TSP筛选）：以步长5m离散化搜索空间，'
            '以TSP路径长度为指标快速筛选，每工位保留Top-3。')
        self.add_para(
            'Phase 2（GA精评）：对候选方案运行GA进行精确评估。')

        self.add_heading('8.3 求解结果', level=2)
        bs = self._r('phase4', 'best_station')
        bx = self._r('phase4', 'best_x', '.0f')
        by = self._r('phase4', 'best_y', '.0f')
        bi = self._r('phase4', 'mksp_improve', '.1f')
        ei = self._r('phase4', 'energy_improve', '.1f')

        if bs is not None:
            self.add_para(
                f'两阶段搜索完成后，最优调整方案为：'
                f'移动工位{bs}至({bx},{by})。')

            self.add_para(
                f'改善效果：makespan改善{bi}%（绝对值），'
                f'总能耗改善{ei}%（绝对值），'
                f'验证了工位布局对AGV调度效率的显著影响。')
        else:
            self.add_para('（运行求解后填入结果。）')

        img = os.path.join(self.result_dir, 'phase4_compare.png')
        self.add_image(img, '图6  布局优化前后对比')

        self.page_break()

    def write_validation(self):
        """模型检验"""
        self.add_heading('九、模型分析与检验', level=1)

        self.add_heading('9.1 收敛性分析', level=2)
        self.add_para(
            'GA收敛曲线表明算法能在有限代数内稳定收敛。'
            'NSGA-II的Pareto前沿随代数逐步前移。')

        self.add_heading('9.2 灵敏度分析', level=2)
        self.add_para(
            '对关键参数进行灵敏度分析：'
            '宽限期δ和容量Q的变化对调度结果的影响在可控范围内，'
            '验证了模型的鲁棒性。'
            '（注：具体分析数据在运行后填入。）')

        self.add_heading('9.3 小规模验证', level=2)
        self.add_para(
            '在较小规模下，将GA结果与枚举法对比，'
            '验证了算法编码和评估逻辑的正确性。')
        self.page_break()

    def write_evaluation(self):
        """模型评价"""
        self.add_heading('十、模型评价', level=1)

        self.add_heading('10.1 模型优点', level=2)
        advantages = [
            '递进式建模思路清晰，各阶段环环相扣。',
            '算法选择基于问题特征分析，针对性强。',
            '多目标优化提供了灵活决策空间。',
            '考虑了实际约束，实用性强。',
        ]
        for i, adv in enumerate(advantages, 1):
            self.add_para(f'（{i}）{adv}')

        self.add_heading('10.2 模型缺点与改进方向', level=2)
        disadvantages = [
            '元启发式算法无法保证全局最优。',
            '未考虑AGV碰撞和路径冲突。',
            '网格搜索步长可能错过更优解。',
            '能耗模型未考虑加减速等非线性因素。',
        ]
        for i, dis in enumerate(disadvantages, 1):
            self.add_para(f'（{i}）{dis}')
        self.page_break()

    def write_references(self):
        """参考文献"""
        self.add_heading('十一、参考文献', level=1)
        refs = [
            '[1] Bellman R. Dynamic programming treatment of the travelling salesman problem[J]. '
            'Journal of the ACM, 1962, 9(1): 61-63.',
            '[2] Held M, Karp R M. A dynamic programming approach to sequencing problems[J]. '
            'Journal of the Society for Industrial and Applied Mathematics, 1962, 10(1): 196-210.',
            '[3] Deb K, Pratap A, Agarwal S, et al. A fast and elitist multiobjective genetic '
            'algorithm: NSGA-II[J]. IEEE Transactions on Evolutionary Computation, 2002, 6(2): 182-197.',
            '[4] Solomon M M. Algorithms for the vehicle routing and scheduling problems with time '
            'window constraints[J]. Operations Research, 1987, 35(2): 254-265.',
            '[5] Mladenović N, Hansen P. Variable neighborhood search[J]. '
            'Computers & Operations Research, 1997, 24(11): 1097-1100.',
            '[6] 司守奎, 孙玺菁. 数学建模算法与应用[M]. 北京: 国防工业出版社, 2011.',
        ]
        for ref in refs:
            self.add_para(ref, indent=False)
        self.page_break()

    def write_appendix(self):
        """附录"""
        self.add_heading('十二、附录', level=1)

        self.add_heading('附录A：程序文件清单', level=2)
        headers = ['文件名', '功能说明']
        rows = [
            ['config.py', '全局配置与参数定义'],
            ['utils.py', '工具函数与可视化'],
            ['tsp_solver.py', 'Phase1: TSP求解'],
            ['vrptw_ga.py', 'Phase2: 增强遗传算法'],
            ['nsga2.py', 'Phase3: NSGA-II+VNS'],
            ['layout_optimizer.py', 'Phase4: 布局优化'],
            ['main.py', '主入口，串联全流程'],
        ]
        self.add_table(headers, rows, caption='表2  程序文件清单')

        self.add_heading('附录B：核心代码', level=2)
        self.add_para(
            '限于篇幅，以下列出关键算法核心逻辑。'
            '完整代码随论文提交。', indent=False)

        self.add_para('Phase1 Held-Karp DP（tsp_solver.py）：', bold=True, indent=False)
        self.add_para(
            'def held_karp():\n'
            '    dp[mask][i] = min(dp[mask xor (1<<i)][j] + dist(j,i) for j in mask)\n'
            '    return min(dp[(1<<n)-1][i] + dist(i,0) for i)',
            indent=False)

        self.add_para('Phase2 GA适应度评估（vrptw_ga.py）：', bold=True, indent=False)
        self.add_para(
            'def evaluate(perm, stations, weights, deadlines):\n'
            '    trips = ffd_pack(perm, weights)\n'
            '    _, makespan, penalty = assign_trips_to_agvs(...)\n'
            '    return -(w1*makespan + w2*penalty), makespan, penalty',
            indent=False)

        self.add_para('Phase3 NSGA-II双目标（nsga2.py）：', bold=True, indent=False)
        self.add_para(
            'def objectives(perm, stations, weights, deadlines):\n'
            '    trips = scan_pack(perm, weights)\n'
            '    energy, agv_times, makespan = compute_trips_energy(...)\n'
            '    return makespan, total_energy',
            indent=False)

        self.add_para('Phase4 两阶段搜索（layout_optimizer.py）：', bold=True, indent=False)
        self.add_para(
            'def optimize(stations, weights, deadlines):\n'
            '    # Phase1: TSP筛选 (648 -> 24)\n'
            '    # Phase2: GA精评 (24 -> best)\n'
            '    return best_station, best_coords, improvements',
            indent=False)

        self.page_break()

    def write_ai_report(self):
        """AI使用报告"""
        self.add_heading('十三、人工智能使用报告', level=1)
        self.add_para('使用AI平台：Claude Code（Anthropic Claude）', indent=False)
        self.doc.add_paragraph()
        self.add_para('AI辅助范围：', bold=True, indent=False)
        self.add_para('（1）算法设计咨询：AI就算法选择和实现细节提供技术建议。')
        self.add_para('（2）代码生成与调试：AI协助生成了算法框架和核心实现，并排查代码错误。')
        self.add_para('（3）论文草稿撰写：AI基于求解结果生成了论文初稿。')
        self.doc.add_paragraph()
        self.add_para('人工审核确认：', bold=True, indent=False)
        self.add_para(
            '所有模型假设、公式推导、求解结果和论文内容均经人工逐项审核确认。'
            'AI仅作为效率工具，最终学术责任由参赛队员承担。')

    # ═══════════════════════════════════════
    # 构建与保存
    # ═══════════════════════════════════════

    def build(self):
        """按顺序构建全部章节"""
        self.write_cover()
        self.write_abstract()
        self.write_restatement()
        self.write_analysis()
        self.write_assumptions()
        self.write_symbols()
        self.write_phase1()
        self.write_phase2()
        self.write_phase3()
        self.write_phase4()
        self.write_validation()
        self.write_evaluation()
        self.write_references()
        self.write_appendix()
        self.write_ai_report()
        return self

    def save(self, output_path):
        """保存文档"""
        self.doc.save(output_path)
        return output_path


# ===== 独立使用入口 =====
if __name__ == '__main__':
    import sys

    # 自动检测 results.json
    candidates = [
        os.path.join(os.path.dirname(__file__), '..', '..', 'examples', 'results', 'results.json'),
        os.path.join(os.getcwd(), 'examples', 'results', 'results.json'),
        os.path.join(os.getcwd(), 'results', 'results.json'),
    ]
    results_path = None
    for p in candidates:
        rp = os.path.abspath(p)
        if os.path.exists(rp):
            results_path = rp
            break

    result_dir = os.path.dirname(results_path) if results_path else os.getcwd()

    config = PaperConfig(title='AGV协同调度与能耗优化问题研究', school_name='XX大学')
    gen = PaperGenerator(config, result_dir=result_dir)

    if results_path:
        print(f'检测到结果文件: {results_path}')
        gen.from_results_json(results_path)
        print('结果已注入论文')
    else:
        print('未检测到 results.json，生成空白示例论文')
        print('  请先运行 examples/main.py 生成结果')

    gen.build()
    out = '示例论文.docx'
    gen.save(out)
    size = os.path.getsize(out) / 1024
    print(f'论文已生成: {out} ({size:.0f} KB)')
